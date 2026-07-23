from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    hs.font.name = 'Calibri'

IMG_DIR = r'C:\Users\ZBOOK STUDIO G5\Documents\GitHub\aas\screenshots'
W = Inches(6.0)

def add_img(name, caption, width=W):
    path = os.path.join(IMG_DIR, name)
    if not os.path.exists(path):
        doc.add_paragraph(f'[Image not found: {name}]').alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r.italic = True

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold:
        r.bold = True
    return p

def section_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CEDRIC MASTERS AUTOS')
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Workshop & Dealership Management System')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('User Manual')
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
r.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Version 1.0  •  July 2026')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

section_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Getting Started — Login & Navigation',
    '3. Administrator — Full System Access',
    '4. CSR / Service Advisor — Front Desk',
    '5. Technician — Workshop Jobs',
    '6. Inventory Manager — Parts & Stock',
    '7. Finance Clerk — Invoices & Payments',
    '8. Manager — Oversight & Approvals',
    '9. Sales Rep — Leads, Orders & Deliveries',
]
for item in toc_items:
    add_para(item, size=12)

section_break()

# ══════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)
add_para(
    'Cedric Masters Autos (CMA) is a web-based workshop and dealership management system. '
    'It helps automotive businesses manage every vehicle that comes through their doors — '
    'from service to sale.'
)
add_para('The system is organized into three core modules:')
add_bullet('After-Sales Service — Customer check-in, technician workflow, spare parts, invoicing')
add_bullet('Vehicle Sales — Inventory tracking, customer leads, sales orders, delivery handover')
add_bullet('Administration — User management, roles, settings')

add_para('')
add_para('Who this manual is for:', bold=True)
add_bullet('Service Advisors (CSR) — Front desk, check-in, customer registration')
add_bullet('Technicians — Diagnose issues, work on jobs, request parts')
add_bullet('Inventory Managers — Manage parts catalogue and stock levels')
add_bullet('Finance Clerks — Generate invoices, record payments')
add_bullet('Managers — Approve invoices, oversee operations')
add_bullet('Sales Reps — Manage leads, create orders, handle deliveries')
add_bullet('Administrators — Full access, manage users, configure settings')

section_break()

# ══════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Getting Started — Login & Navigation', level=1)

doc.add_heading('2.1 Accessing the System', level=2)
add_para('Open your web browser and navigate to the CMA system URL. You will see the login screen.')

add_img('01-login.png', 'Login Screen — Enter your email and password to sign in.')

doc.add_heading('2.2 Signing In', level=2)
add_para('Enter your registered email address and password, then click "Sign in".')
add_para('If you do not have an account, click "Create an account instead" to register. '
         'An administrator must assign your role before you can access the system.')

doc.add_heading('2.3 Dashboard Overview', level=2)
add_para('After signing in, the Dashboard displays key performance indicators, active jobs, '
         'recent check-ins, and quick links to your most-used sections.')

add_img('02-dashboard.png', 'Dashboard — KPI cards, active jobs table, finance snapshot, and quick links.')

doc.add_heading('2.4 Navigation', level=2)
add_para('The sidebar on the left provides access to all sections. Sections are grouped:')
add_bullet('General — Dashboard, Appointments, Customers, Jobs')
add_bullet('Sales — Inventory, Leads, Orders')
add_bullet('Operations — Customer Vehicles, Parts, Finance, User Management')
add_para('Your role determines which menu items you can see. The top bar shows your current '
         'location with breadcrumbs and provides dark mode toggle, notifications, and profile access.')

section_break()

# ══════════════════════════════════════════════════════════
# 3. ADMINISTRATOR
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Administrator — Full System Access', level=1)
add_para('Administrators have unrestricted access to every feature. Your primary responsibilities '
         'include managing users, configuring system settings, and overseeing all operations.')

doc.add_heading('3.1 User Management', level=2)
add_para('Navigate to Operations > User Management to view, add, and manage staff accounts. '
         'Each user can be assigned one of seven roles that control their access level.')

add_img('14-user-management.png', 'User Management — View all staff, assign roles, manage access.')

doc.add_heading('3.2 Role Assignments', level=2)
add_para('The seven roles available are:')
add_bullet('Admin — Full system access')
add_bullet('CSR — Front desk, appointments, check-in')
add_bullet('Technician — Jobs, diagnosis, parts requests')
add_bullet('Inventory Manager — Parts catalogue, stock')
add_bullet('Finance Clerk — Invoices, payments')
add_bullet('Manager — Oversight, approvals')
add_bullet('Sales Rep — Leads, sales orders, deliveries')

doc.add_heading('3.3 Finance Configuration', level=2)
add_para('Configure VAT rate and labour types under Operations > Finance. These settings '
         'affect invoice calculations across the system.')

add_img('10-finance.png', 'Finance Settings — Configure VAT rate (default 7.5%) and labour type flat fees.')

doc.add_heading('3.4 Parts Catalogue Management', level=2)
add_para('The Parts Catalogue under Operations > Parts allows full CRUD operations on spare parts, '
         'including Excel import for bulk data loading, stock adjustments, and low-stock alerts.')

add_img('09-parts.png', 'Parts Catalogue — Manage parts, stock levels, and import from Excel.')

doc.add_heading('3.5 Appointments', level=2)
add_para('View and manage all customer appointments under General > Appointments. '
         'Appointments can be viewed by day, week, or month ranges.')

add_img('03-appointments.png', 'Appointments — Range view with day/week/month presets.')

section_break()

# ══════════════════════════════════════════════════════════
# 4. CSR / SERVICE ADVISOR
# ══════════════════════════════════════════════════════════
doc.add_heading('4. CSR / Service Advisor — Front Desk', level=1)
add_para('As a Service Advisor (CSR), you are the first point of contact for customers. '
         'Your workflow covers customer registration, vehicle check-in, appointment booking, '
         'and managing the customer database.')

doc.add_heading('4.1 Customer Management', level=2)
add_para('Navigate to General > Customers to view the customer directory. '
         'You can search by name or phone number, view repair history, and add new customers.')
add_para('The one-stop creation flow lets you add a customer, their vehicle, and a new job '
         'in a single process.')

add_img('04-customers.png', 'Customers — Searchable directory with repair history per customer.')

doc.add_heading('4.2 Vehicle Check-In', level=2)
add_para('Click "Check In Vehicle" from the Dashboard or navigate to General > Jobs > Check In. '
         'The 3-step flow captures vehicle details, customer information, and the reported issue.')

add_img('07-checkin.png', 'Check-In Flow — 3-step process to register a new service job.')

doc.add_heading('4.3 Appointment Booking', level=2)
add_para('Under General > Appointments, you can book appointments for customers. '
         'Vehicle details (make, model, plate) and complaint are mandatory fields.')

add_img('03-appointments.png', 'Appointments — Book and manage customer appointments with day navigation.')

doc.add_heading('4.4 Customer Vehicles', level=2)
add_para('View all vehicles registered in the system under Operations > Customer Vehicles. '
         'Each vehicle is linked to its owner and service history.')

add_img('08-vehicles.png', 'Customer Vehicles — Directory of all vehicles with owner details.')

section_break()

# ══════════════════════════════════════════════════════════
# 5. TECHNICIAN
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Technician — Workshop Jobs', level=1)
add_para('Technicians are responsible for diagnosing issues, performing repairs, and requesting '
         'parts. Your workflow is centered around the Jobs section.')

doc.add_heading('5.1 Viewing Assigned Jobs', level=2)
add_para('Navigate to General > Jobs to see the list of all jobs in the workshop. '
         'Filter by status to focus on your assigned work.')

add_img('05-jobs.png', 'Jobs List — All workshop jobs with status, vehicle, and customer details.')

doc.add_heading('5.2 Job Detail & Diagnosis', level=2)
add_para('Click on any job to open its detail page. The status stepper shows the current stage '
         'of the job. You can record your diagnosis, add job items (parts and labour), '
         'and transition the job through statuses.')

add_img('06-job-detail.png', 'Job Detail — Status stepper, diagnosis, parts requests, and job items.')

doc.add_heading('5.3 Parts Requests', level=2)
add_para('When parts are needed, create a parts request from the job detail page. '
         'The Inventory Manager will approve and dispatch the parts. Dispatched parts '
         'are automatically added to the job items and invoice.')

add_para('', bold=False)
add_para('Note: Technicians cannot view or modify invoice information. '
         'Parts requests are restricted to Technicians and Administrators.')

doc.add_heading('5.4 Printable Job Card', level=2)
add_para('Each job has a printable job card accessible via the Print button. '
         'This card shows customer info, vehicle details, diagnosis, and job items '
         'in a print-friendly format.')

section_break()

# ══════════════════════════════════════════════════════════
# 6. INVENTORY MANAGER
# ══════════════════════════════════════════════════════════
doc.add_heading('6. Inventory Manager — Parts & Stock', level=1)
add_para('Inventory Managers are responsible for the spare parts catalogue, stock levels, '
         'and fulfilling technician parts requests.')

doc.add_heading('6.1 Parts Catalogue', level=2)
add_para('Under Operations > Parts, view the complete parts catalogue. Each part has a unique '
         'code, OEM number, cost price, selling price, and stock quantity.')

add_img('09-parts.png', 'Parts Catalogue — Search, filter, edit parts, and view low-stock alerts.')

doc.add_heading('6.2 Stock Management', level=2)
add_para('Use the stock adjustment feature to record stock movements (in/out/adjust). '
         'Low-stock alerts appear when quantity falls below the reorder level. '
         'All stock movements are tracked in an audit trail.')

doc.add_heading('6.3 Excel Import', level=2)
add_para('Bulk-import parts using the Excel/CSV import feature. The system validates the data '
         'and adds parts to the catalogue in one operation.')

doc.add_heading('6.4 Parts Request Fulfillment', level=2)
add_para('When a Technician requests parts, you will receive the request in the job detail. '
         'Review stock availability, approve or reject, and dispatch the parts. '
         'Dispatched parts are automatically deducted from inventory and added to the job.')

add_para('')
add_para('Note: A confirmation modal shows stock status before finalizing dispatch. '
         'You can also reverse a dispatch to return items to inventory.')

section_break()

# ══════════════════════════════════════════════════════════
# 7. FINANCE CLERK
# ══════════════════════════════════════════════════════════
doc.add_heading('7. Finance Clerk — Invoices & Payments', level=1)
add_para('Finance Clerks handle invoice generation, approval workflow, and payment recording.')

doc.add_heading('7.1 Finance Settings', level=2)
add_para('Under Operations > Finance, view and manage the VAT rate and labour type configurations. '
         'These settings are used when calculating invoices.')

add_img('10-finance.png', 'Finance Dashboard — Configure VAT, manage labour types with fixed prices.')

doc.add_heading('7.2 Invoice Generation', level=2)
add_para('From any job detail page, generate an invoice. The invoice auto-calculates totals from '
         'parts and labour line items, applies VAT at the configured rate, and locks line items '
         'at the time of generation.')

doc.add_heading('7.3 Manager Approval', level=2)
add_para('Generated invoices require Manager approval before they are finalized. '
         'The invoice status shows its current stage in the approval workflow.')

doc.add_heading('7.4 Payment Recording', level=2)
add_para('Record payments against invoices using the payment panel on the job detail page. '
         'Supported methods:')
add_bullet('Cash')
add_bullet('Bank Transfer')
add_bullet('Card / POS')
add_para('Partial payments are supported with automatic balance tracking. '
         'When fully paid, the job auto-transitions to "Paid" status.')

doc.add_heading('7.5 Invoice Regeneration', level=2)
add_para('If job items change after invoice generation, use the "Regenerate" button '
         'to re-sync the invoice with the latest items.')

doc.add_heading('7.6 Printable Invoice', level=2)
add_para('Each invoice can be printed in a clean, print-friendly format via the Print button.')

section_break()

# ══════════════════════════════════════════════════════════
# 8. MANAGER
# ══════════════════════════════════════════════════════════
doc.add_heading('8. Manager — Oversight & Approvals', level=1)
add_para('Managers have broad access across all modules for oversight. Your key actions include '
         'approving invoices, monitoring workshop progress, and managing customer relationships.')

doc.add_heading('8.1 Dashboard & Monitoring', level=2)
add_para('The Dashboard gives you a real-time view of open jobs, in-progress work, '
         'vehicles ready for pickup, and customer metrics for the month.')

add_img('02-dashboard.png', 'Dashboard — KPI cards, active jobs, recent check-ins for full oversight.')

doc.add_heading('8.2 Invoice Approval', level=2)
add_para('Review and approve invoices generated by Finance Clerks. Navigate to the job detail '
         'to see pending invoices and approve them with a single click.')

doc.add_heading('8.3 Job Completion', level=2)
add_para('Managers (and Admins) can mark jobs as "Completed" after work is finished. '
         'This is a restricted action — Technicians and CSRs cannot complete jobs.')

doc.add_heading('8.4 Customer & Vehicle Viewing', level=2)
add_para('Access the full customer database and vehicle directory to review service history '
         'and customer details.')

doc.add_heading('8.5 Sales Overview', level=2)
add_para('View sales inventory, leads pipeline, and sales orders. Monitor lead stages '
         'and review sales team performance.')

add_img('11-sales-inventory.png', 'Sales Inventory — Vehicle stock tracking with sales status.')

add_img('12-sales-leads.png', 'Sales Leads — Pipeline view with lead stages and follow-up tracking.')

add_img('13-sales-orders.png', 'Sales Orders — Order management with deposit and balance tracking.')

section_break()

# ══════════════════════════════════════════════════════════
# 9. SALES REP
# ══════════════════════════════════════════════════════════
doc.add_heading('9. Sales Rep — Leads, Orders & Deliveries', level=1)
add_para('Sales Representatives manage the full sales cycle from lead generation to vehicle delivery.')

doc.add_heading('9.1 Vehicle Inventory', level=2)
add_para('Under Sales > Inventory, view all vehicles available for sale. Track stock quantities, '
         'reorder levels, and adjust stock as needed. Vehicles have status: In Stock / Reserved / Sold.')

add_img('11-sales-inventory.png', 'Sales Inventory — Vehicle list with make, model, year, status, and stock levels.')

doc.add_heading('9.2 Lead Management', level=2)
add_para('Navigate to Sales > Leads to manage customer enquiries. Track leads through stages:')
add_bullet('New')
add_bullet('Contacted')
add_bullet('Qualified')
add_bullet('Won / Lost')
add_para('Log follow-up notes and set reminders directly on the lead detail page.')

add_img('12-sales-leads.png', 'Leads — Searchable list with stage filters and create lead action.')

doc.add_heading('9.3 Creating Sales Orders', level=2)
add_para('From the lead detail page or Sales > Orders, create a new sales order. '
         'Select the vehicle, set the agreed price and deposit amount. '
         'The vehicle is automatically reserved.')

add_img('13-sales-orders.png', 'Sales Orders — Order list with deposit and balance tracking.')

doc.add_heading('9.4 Recording Payments', level=2)
add_para('On the order detail page, record customer payments. Track deposits and remaining balances. '
         'Partial payments are supported.')

doc.add_heading('9.5 Delivery Handover', level=2)
add_para('When a sale is complete, use the Delivery Handover checklist. '
         'The checklist covers:')
add_bullet('Keys handover')
add_bullet('Owner\'s manual')
add_bullet('Toolkit')
add_bullet('Vehicle inspection')
add_para('The handover is recorded against the sales rep.')

section_break()

# ══════════════════════════════════════════════════════════
# TIPS
# ══════════════════════════════════════════════════════════
doc.add_heading('Quick Tips', level=1)
add_bullet('Use the sidebar to navigate between sections. Your role determines visible menu items.')
add_bullet('The search bar at the top adapts its placeholder based on your current page.')
add_bullet('Toggle dark mode using the sun/moon icon in the top bar.')
add_bullet('Your profile (top-right corner) shows your name, email, and role. Sign out from here.')
add_bullet('All actions are logged in the audit trail for accountability.')
add_bullet('Use the Print button on job cards and invoices for physical documentation.')

doc.add_paragraph()
add_para('— End of Manual —', bold=True)

# ── SAVE ──────────────────────────────────────────────────
output_path = r'C:\Users\ZBOOK STUDIO G5\Documents\GitHub\aas\Cedric_Masters_Autos_User_Manual.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
